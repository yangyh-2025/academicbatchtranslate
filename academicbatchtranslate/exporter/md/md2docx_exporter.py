# SPDX-License-Identifier: MPL-2.0
import base64
import io
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict

import docx.opc.constants
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor

from academicbatchtranslate.exporter.base import Exporter, ExporterConfig
from academicbatchtranslate.exporter.md.types import MD2DocxEngineType
from academicbatchtranslate.ir.document import Document
from academicbatchtranslate.ir.markdown_document import MarkdownDocument
from academicbatchtranslate.logger import global_logger


@dataclass(kw_only=True)
class MD2DocxExporterConfig(ExporterConfig):
    engine: MD2DocxEngineType = "auto"


def is_pandoc_available() -> bool:
    """检测pandoc是否可用"""
    return shutil.which("pandoc") is not None


def _md_to_docx_via_pandoc(md_content: str, logger=global_logger) -> bytes:
    """使用pandoc将markdown转换为docx"""
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        md_file = tmpdir / "input.md"
        docx_file = tmpdir / "output.docx"

        # 写入markdown文件（使用utf-8-sig确保中文兼容）
        md_file.write_text(md_content, encoding="utf-8-sig")

        # 执行pandoc转换
        try:
            result = subprocess.run(
                ["pandoc", str(md_file), "-o", str(docx_file)],
                capture_output=True,
                text=True,
                check=True
            )
            logger.info(f"Pandoc转换成功: {result.stdout}")
        except subprocess.CalledProcessError as e:
            logger.error(f"Pandoc转换失败: {e.stderr}")
            raise RuntimeError(f"Pandoc转换失败: {e.stderr}")

        return docx_file.read_bytes()


# =============================================================================
# 增强版纯Python转换核心逻辑
# =============================================================================

def _add_hyperlink(paragraph, url: str, text: str, color="0563C1", underline=True):
    """
    通过底层OXML在段落中插入可点击的超链接
    """
    part = paragraph.part
    try:
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    except Exception:
        paragraph.add_run(text)
        return

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_image_run(paragraph, img_info: dict, width_inch: float = 5.8):
    """在段落中添加图片Run"""
    try:
        mime = img_info.get('mime', '')
        ext = 'jpg'
        if 'png' in mime:
            ext = 'png'
        elif 'gif' in mime:
            ext = 'gif'
        elif 'bmp' in mime:
            ext = 'bmp'

        if 'svg' in mime: return

        with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as tmp:
            tmp.write(img_info['data'])
            tmp_path = tmp.name

        try:
            run = paragraph.add_run()
            run.add_picture(tmp_path, width=Inches(width_inch))
        finally:
            Path(tmp_path).unlink()
    except Exception:
        paragraph.add_run(f"[Image Error: {img_info.get('alt', 'unk')}]")


def _process_inline_markdown(paragraph, text: str, images: Dict[int, Dict]):
    """
    行内解析器：支持 粗体、斜体、代码、链接、图片 的混合解析
    """
    patterns = [
        r'(<!--IMG_PLACEHOLDER_\d+-->)',
        r'(`[^`]+`)',
        r'(!\[[^\]]*\]\([^\)]+\))',
        r'(\[[^\]]+\]\([^\)]+\))',
        r'(\*\*[^\*]+\*\*)',
        r'(__[^_]+__)',
        r'(\*[^\*]+\*)',
        r'(_[^_]+_)',
    ]

    master_pattern = re.compile('|'.join(patterns))
    parts = master_pattern.split(text)

    for part in parts:
        if not part: continue

        if part.startswith('<!--IMG_PLACEHOLDER_'):
            m = re.match(r'<!--IMG_PLACEHOLDER_(\d+)-->', part)
            if m:
                idx = int(m.group(1))
                if idx in images:
                    _add_image_run(paragraph, images[idx])
            continue

        if part.startswith('`') and part.endswith('`') and len(part) > 1:
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(199, 37, 78)
            continue

        if part.startswith('![') and ']' in part and '(' in part and part.endswith(')'):
            m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', part)
            if m:
                alt = m.group(1)
                paragraph.add_run(f"[Image: {alt}]").italic = True
            continue

        if part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                link_text, link_url = m.group(1), m.group(2)
                _add_hyperlink(paragraph, link_url, link_text)
            else:
                paragraph.add_run(part)
            continue

        if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
            continue

        if (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
            continue

        paragraph.add_run(part)


def _process_block_code(doc, code_lines: List[str]):
    """处理代码块"""
    if not code_lines: return
    text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(36, 41, 46)


def _render_table_data(doc, rows: List[List[str]]):
    """
    通用表格渲染函数（用于 Markdown 表格和 HTML 表格）
    """
    if not rows: return

    # 确定列数 (取最大列数)
    col_count = max(len(row) for row in rows)
    if col_count == 0: return

    # 创建表格
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'

    # 填充数据
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx].cells
        for c_idx, text in enumerate(row_data):
            if c_idx < len(row_cells):
                # 如果是第一行，默认为表头
                if r_idx == 0:
                    p = row_cells[c_idx].paragraphs[0]
                    run = p.add_run(str(text))
                    run.bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    row_cells[c_idx].text = str(text)

    doc.add_paragraph()  # 表格后空行


def _process_markdown_table(doc, table_lines: List[str]):
    """处理Markdown语法的表格"""
    if len(table_lines) < 2: return
    rows = []
    for line in table_lines:
        content = line.strip().strip('|')
        cells = [c.strip() for c in content.split('|')]
        rows.append(cells)

    if len(rows) < 2: return

    # 校验第二行是否为分割线
    separator = rows[1]
    is_valid = True
    for cell in separator:
        if not re.match(r'^[\s\-\:]+$', cell):
            is_valid = False;
            break
    if not is_valid:
        for line in table_lines: doc.add_paragraph(line)
        return

    # 移除分割线行
    data_rows = [rows[0]] + rows[2:]
    _render_table_data(doc, data_rows)


def _md_to_docx_via_python(md_content: str, logger=global_logger) -> bytes:
    """
    纯Python Markdown -> Docx 转换器
    支持：Markdown语法 + 原生HTML表格(<table>)
    """
    if DocxDocument is None:
        raise RuntimeError("依赖缺失: 未安装 python-docx，无法执行转换。")

    doc = DocxDocument()

    # 0. 全局样式设置
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.space_after = Pt(8)

    # 1. 图片预处理
    images = {}
    img_idx = 0

    def extract_base64(match):
        nonlocal img_idx
        alt = match.group(1)
        data_uri = match.group(2)
        try:
            header, b64 = data_uri.split(',', 1)
            mime = header.split(';')[0].replace('data:', '')
            data = base64.b64decode(b64)
            key = img_idx
            images[key] = {'alt': alt, 'mime': mime, 'data': data}
            placeholder = f"<!--IMG_PLACEHOLDER_{key}-->"
            img_idx += 1
            return placeholder
        except Exception:
            return match.group(0)

    md_content = re.sub(r'!\[([^\]]*)\]\((data:image/[^)]+)\)', extract_base64, md_content)

    # 2. HTML 表格预处理 (利用 BeautifulSoup)
    html_tables = {}
    html_table_idx = 0

    if BeautifulSoup:
        def extract_html_table(match):
            nonlocal html_table_idx
            html_str = match.group(0)
            try:
                soup = BeautifulSoup(html_str, 'html.parser')
                rows_data = []
                for tr in soup.find_all('tr'):
                    # 同时支持 th 和 td
                    cells = [cell.get_text(strip=True) for cell in tr.find_all(['td', 'th'])]
                    if cells:
                        rows_data.append(cells)

                if rows_data:
                    key = html_table_idx
                    html_tables[key] = rows_data
                    placeholder = f"\n<!--HTML_TABLE_PLACEHOLDER_{key}-->\n"
                    html_table_idx += 1
                    return placeholder
                return html_str  # 解析失败则保留原样
            except Exception:
                return html_str

        # 匹配 <table>...</table>，使用 DOTALL 模式匹配跨行
        md_content = re.sub(r'<table[^>]*>.*?</table>', extract_html_table, md_content, flags=re.DOTALL | re.IGNORECASE)

    # 3. 逐行解析
    lines = md_content.split('\n')
    n = len(lines)
    i = 0

    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        # A. 代码块
        if stripped.startswith('```') or stripped.startswith('~~~'):
            if in_code_block:
                _process_block_code(doc, code_buffer)
                in_code_block = False
                code_buffer = []
            else:
                if in_table:
                    _process_markdown_table(doc, table_buffer)
                    in_table = False
                    table_buffer = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # B. Markdown 表格
        is_md_table_row = stripped.startswith('|') or (stripped.endswith('|') and '|' in stripped)
        if is_md_table_row:
            if not in_table:
                if i + 1 < n:
                    next_line = lines[i + 1].strip()
                    if set(next_line) & {'-', '|'} and re.match(r'^\|?[\s\-:|]+\|?$', next_line):
                        in_table = True
                        table_buffer = [line]
                        i += 1
                        continue
            else:
                table_buffer.append(line)
                i += 1
                continue

        if in_table:
            _process_markdown_table(doc, table_buffer)
            in_table = False
            table_buffer = []
            if not is_md_table_row:
                pass
            else:
                i += 1;
                continue

        # C. HTML 表格占位符处理
        if stripped.startswith('<!--HTML_TABLE_PLACEHOLDER_') and stripped.endswith('-->'):
            m = re.match(r'<!--HTML_TABLE_PLACEHOLDER_(\d+)-->', stripped)
            if m:
                idx = int(m.group(1))
                if idx in html_tables:
                    _render_table_data(doc, html_tables[idx])
                i += 1
                continue

        # D. 普通元素
        if stripped.startswith('#'):
            level = 0
            for char in stripped:
                if char == '#':
                    level += 1
                else:
                    break
            if 0 < level <= 6:
                text = stripped[level:].strip()
                heading = doc.add_heading(level=level)
                _process_inline_markdown(heading, text, images)
                i += 1
                continue

        if re.match(r'^(\*{3,}|-{3,}|_{3,})$', stripped):
            p = doc.add_paragraph()
            run = p.add_run("________________________________________")
            run.font.color.rgb = RGBColor(220, 220, 220)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _process_inline_markdown(p, text, images)
            for run in p.runs:
                run.font.color.rgb = RGBColor(120, 120, 120)
            i += 1
            continue

        ul_match = re.match(r'^(\s*)([\*\-\+])\s+(.*)', line)
        ol_match = re.match(r'^(\s*)(\d+\.)\s+(.*)', line)
        if ul_match or ol_match:
            if ul_match:
                indent, _, text = ul_match.groups()
                style = 'List Bullet'
            else:
                indent, _, text = ol_match.groups()
                style = 'List Number'
            p = doc.add_paragraph(style=style)
            level = len(indent.replace('\t', '  ')) // 2
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            _process_inline_markdown(p, text, images)
            i += 1
            continue

        if stripped.startswith('<!--IMG_PLACEHOLDER_') and stripped.endswith('-->'):
            if stripped.count('<!--IMG_PLACEHOLDER_') == 1 and len(stripped) < 40:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                _process_inline_markdown(p, stripped, images)
                i += 1
                continue

        if stripped:
            p = doc.add_paragraph()
            _process_inline_markdown(p, stripped, images)

        i += 1

    if in_table and table_buffer:
        _process_markdown_table(doc, table_buffer)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class MD2DocxExporter(Exporter):
    def __init__(self, config: MD2DocxExporterConfig | None = None):
        config = config or MD2DocxExporterConfig()
        super().__init__(config=config)
        self.engine = config.engine
        self.logger = config.logger if hasattr(config, 'logger') and config.logger else global_logger

    def export(self, document: MarkdownDocument) -> Document:
        md_content = document.content.decode("utf-8")

        engine = self.engine
        if engine == "auto":
            engine = "pandoc" if is_pandoc_available() else "python"

        if engine == "pandoc":
            if not is_pandoc_available():
                self.logger.warning("Pandoc不可用，回退到纯Python模式")
                docx_bytes = _md_to_docx_via_python(md_content, self.logger)
            else:
                docx_bytes = _md_to_docx_via_pandoc(md_content, self.logger)
        else:
            docx_bytes = _md_to_docx_via_python(md_content, self.logger)

        return Document.from_bytes(docx_bytes, suffix=".docx", stem=document.stem)
